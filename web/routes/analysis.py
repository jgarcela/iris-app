# web/analysis.py
import json
import re
from flask import Blueprint, render_template, request, jsonify, session, abort, current_app, g
import requests
from flask_babel import get_locale
import configparser
import ast
from datetime import datetime
from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import io

from web.utils.logger import logger
from web.utils.decorators import login_required
import database.db as db
from bson import ObjectId

from flask import render_template, make_response, current_app, request
from weasyprint import HTML

# ----------------- BLUEPRINT -----------------
bp = Blueprint(
    'analysis',
    __name__,
    url_prefix='/analysis'
)

# ----------------- CONFIG -----------------
config = configparser.ConfigParser()
config.read('config.ini')

WEB_HOST = config['WEB']['HOST']
WEB_PORT = config['WEB']['PORT']
DEBUG = config['WEB'].getboolean('DEBUG')
API_HOST = config['API']['HOST']
API_PORT = config['API']['PORT']
API_HEADERS_str = config['API']['HEADERS']
API_HEADERS = ast.literal_eval(API_HEADERS_str)

ENDPOINT_ANALYSIS = config['API']['ENDPOINT_ANALYSIS']
ENDPOINT_ANALYSIS_ANALYZE = config['API']['ENDPOINT_ANALYSIS_ANALYZE']
ENDPOINT_ANALYSIS_EDITS = config['API']['ENDPOINT_ANALYSIS_EDITS']

ENDPOINT_DATA = config['API']['ENDPOINT_DATA']
ENDPOINT_DATA_GET_DOCUMENT = config['API']['ENDPOINT_DATA_GET_DOCUMENT']

# ----------------- VARIABLES -----------------
HIGHLIGHT_COLOR_MAP = ast.literal_eval(config['VARIABLES']['HIGHLIGHT_COLOR_MAP'])
CONTENIDO_GENERAL_VARIABLES = ast.literal_eval(config['CONTENIDO_GENERAL']['VARIABLES'])
LENGUAJE_VARIABLES = ast.literal_eval(config['LENGUAJE']['VARIABLES'])
FUENTES_VARIABLES = ast.literal_eval(config['FUENTES']['VARIABLES'])

# ----------------- URLs -----------------
URL_API_ENDPOINT_ANALYSIS_ANALYZE = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_ANALYSIS}/{ENDPOINT_ANALYSIS_ANALYZE}"
URL_API_ENDPOINT_ANALYSIS_EDITS = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_ANALYSIS}/{ENDPOINT_ANALYSIS_EDITS}"
URL_API_ENDPOINT_ANALYSIS_SAVE_ANNOTATIONS = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_ANALYSIS}/save_annotations"

URL_API_ENDPOINT_DATA = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_DATA}"
URL_API_ENDPOINT_DATA_GET_DOCUMENT = f"{URL_API_ENDPOINT_DATA}/{ENDPOINT_DATA_GET_DOCUMENT}"

#  ----------------- HELPER FUNCTIONS -----------------
def create_analysis_document(api_data, text, title, authors, url, model, analysis_mode):
    """Create analysis document for database storage"""
    timestamp = datetime.now().isoformat()
    
    # Extract analysis data from API response
    analysis = api_data.get('analysis', {})
    original = analysis.get('original', {})
    edited = analysis.get('edited', {})
    
    # Convert API format to database format
    annotations = []
    
    # Process contenido_general
    if original.get('contenido_general'):
        cg_data = original['contenido_general']
        for variable, value in cg_data.items():
            if value is not None:
                annotations.append({
                    'id': f"{timestamp}_{variable}",
                    'text': text[:100] + "..." if len(text) > 100 else text,  # Sample text
                    'category': 'contenido_general',
                    'variable': variable,
                    'value': str(value),
                    'timestamp': timestamp,
                    'source': 'ai_analysis'
                })
    
    # Process lenguaje
    if original.get('lenguaje'):
        lang_data = original['lenguaje']
        for variable, value in lang_data.items():
            if value is not None:
                annotations.append({
                    'id': f"{timestamp}_{variable}",
                    'text': text[:100] + "..." if len(text) > 100 else text,
                    'category': 'lenguaje',
                    'variable': variable,
                    'value': str(value),
                    'timestamp': timestamp,
                    'source': 'ai_analysis'
                })
    
    # Process fuentes
    if original.get('fuentes'):
        fuentes_data = original['fuentes']
        if 'fuentes' in fuentes_data:
            for i, fuente in enumerate(fuentes_data['fuentes']):
                for variable, value in fuente.items():
                    if value is not None:
                        annotations.append({
                            'id': f"{timestamp}_fuente_{i}_{variable}",
                            'text': fuente.get('declaracion_fuente', ''),
                            'category': 'fuentes',
                            'variable': variable,
                            'value': str(value),
                            'timestamp': timestamp,
                            'source': 'ai_analysis'
                        })
    
    # Determine protagonist analysis
    protagonist_analysis = analyze_protagonist_from_annotations(annotations)
    
    # Create document
    analysis_doc = {
        'user_id': session.get('user_id'),
        'text': text,
        'title': title,
        'authors': authors,
        'url': url,
        'analysis_mode': analysis_mode,
        'model': model,
        'status': 'completed',
        'annotations': annotations,
        'ai_analysis': original,
        'human_edits': edited,
        'protagonist_analysis': protagonist_analysis,
        'created_at': timestamp,
        'updated_at': timestamp,
        'metadata': {
            'total_annotations': len(annotations),
            'categories': list(set([ann.get('category', '') for ann in annotations])),
            'variables': list(set([ann.get('variable', '') for ann in annotations])),
            'ai_model': model,
            'has_human_edits': bool(edited and any(edited.values()))
        }
    }
    
    return analysis_doc

def analyze_protagonist_from_annotations(annotations):
    """Analyze protagonist from annotations"""
    gender_annotations = [ann for ann in annotations 
                         if ann.get('variable') in ['genero_personas_mencionadas', 'genero_nombre_propio_titular']]
    
    if not gender_annotations:
        return {'protagonist': 'No identificado', 'confidence': 0}
    
    male_count = sum(1 for ann in gender_annotations if ann.get('value') == '2')
    female_count = sum(1 for ann in gender_annotations if ann.get('value') == '3')
    mixed_count = sum(1 for ann in gender_annotations if ann.get('value') == '4')
    
    if male_count > female_count and male_count > mixed_count:
        return {'protagonist': 'Masculino', 'confidence': male_count / len(gender_annotations)}
    elif female_count > male_count and female_count > mixed_count:
        return {'protagonist': 'Femenino', 'confidence': female_count / len(gender_annotations)}
    elif mixed_count > 0:
        return {'protagonist': 'Mixto', 'confidence': mixed_count / len(gender_annotations)}
    
    return {'protagonist': 'No identificado', 'confidence': 0}

def get_user_info():
    """Get current user information from context processor"""
    # The current_user is injected by the context processor in __init__.py
    # We need to get it from the template context
    from flask import has_request_context
    if has_request_context():
        # Try to get user info from the request context
        # This is a workaround since current_user is only available in templates
        token = request.cookies.get('access_token_cookie')
        if token:
            try:
                headers = {'Authorization': f'Bearer {token}'}
                resp = requests.get(f"http://{API_HOST}:{API_PORT}/auth/me", headers=headers, timeout=2)
                if resp.ok:
                    user = resp.json().get('user')
                    return user
            except Exception:
                pass
    return None

#  ----------------- ENDPOINTS -----------------
@bp.route('/generate_report/<doc_id>', methods=['GET'])
@login_required
def generate_report(doc_id):
    # 1) Llamada al API para obtener el documento
    url = f"{URL_API_ENDPOINT_DATA_GET_DOCUMENT}/{doc_id}"
    resp = requests.get(url)
    if resp.status_code == 404:
        abort(404, description="Documento no encontrado")
    if resp.status_code != 200:
        abort(resp.status_code, description="Error al obtener el documento")
    data = resp.json()

    # 2) Listado de secciones para el índice
    sections = list(data.get('highlight', {}).get('original', {}).keys())

    # 3) Renderizar HTML con fecha y secciones
    html = render_template(
        'analysis/report.html',
        data=data,
        highlight_map=HIGHLIGHT_COLOR_MAP,
        generation_date=datetime.now(),
        sections=sections
    )

    # 4) Generar PDF
    pdf = HTML(string=html, base_url=request.url_root).write_pdf()

    # 5) Devolver PDF
    response = make_response(pdf)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=Informe_{doc_id}.pdf'
    return response


@bp.route('/generate_report_word/<doc_id>', methods=['GET'])
@login_required
def generate_report_word(doc_id):
    # 1. Obtener datos
    url = f"{URL_API_ENDPOINT_DATA_GET_DOCUMENT}/{doc_id}"
    resp = requests.get(url)
    if resp.status_code == 404:
        abort(404, description="Documento no encontrado")
    if resp.status_code != 200:
        abort(resp.status_code, description="Error al obtener el documento")
    data = resp.json()

    # 2. Colores válidos
    highlight_map_docx = {
        "color-1": WD_COLOR_INDEX.YELLOW,
        "color-2": WD_COLOR_INDEX.BRIGHT_GREEN,
        "color-3": WD_COLOR_INDEX.TURQUOISE,
        "color-4": WD_COLOR_INDEX.PINK,
        "color-5": WD_COLOR_INDEX.BLUE,
        "color-6": WD_COLOR_INDEX.RED,
        "color-7": WD_COLOR_INDEX.GRAY_25,
        "color-8": WD_COLOR_INDEX.DARK_YELLOW,
        "color-9": WD_COLOR_INDEX.GREEN,
        "color-10": WD_COLOR_INDEX.GRAY_50,
        "color-11": WD_COLOR_INDEX.VIOLET,
        "color-12": WD_COLOR_INDEX.TEAL,
        "color-13": WD_COLOR_INDEX.PINK,
        "color-14": WD_COLOR_INDEX.DARK_RED,
        "color-15": WD_COLOR_INDEX.DARK_BLUE,
    }

    # 3. Documento
    doc = Document()
    doc.add_heading("Informe de análisis", 0)

    if data.get("title"):
        doc.add_heading(data["title"], level=1)

    if data.get("authors"):
        authors = data["authors"] if isinstance(data["authors"], str) else ", ".join(data["authors"])
        doc.add_paragraph(f"Autoría: {authors}")

    if data.get("url"):
        doc.add_paragraph(f"Fuente: {data['url']}")

    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_page_break()

    # 4. Índice
    doc.add_heading("Índice", level=1)
    for i, section in enumerate(data.get("highlight", {}).get("original", {}).keys(), 1):
        doc.add_paragraph(f"{i}. {section.replace('_', ' ').capitalize()}", style='List Number')
    doc.add_page_break()

    # 5. Secciones
    for i, (section, html) in enumerate(data["highlight"]["original"].items(), 1):
        doc.add_heading(f"{i}. {section.replace('_',' ').capitalize()}", level=2)
        soup = BeautifulSoup(html, "html.parser")
        paragraph = doc.add_paragraph()

        for elem in soup.recursiveChildGenerator():
            if isinstance(elem, NavigableString):
                if isinstance(elem.parent, Tag) and elem.parent.name == "mark":
                    continue  # ya se procesa desde el <mark>
                if elem.strip():
                    paragraph.add_run(str(elem))

            elif isinstance(elem, Tag) and elem.name == "mark":
                text = elem.get_text()
                cls = elem.get("class", [])
                highlight = highlight_map_docx.get(cls[0]) if cls else None
                run = paragraph.add_run(text)
                if highlight:
                    run.font.highlight_color = highlight

            elif isinstance(elem, Tag) and elem.name == "br":
                paragraph = doc.add_paragraph()

    # 6. Devolver Word
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    response = make_response(doc_io.read())
    response.headers["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    response.headers["Content-Disposition"] = f"attachment; filename=Informe_{doc_id}.docx"
    return response


@bp.route('/create_analysis', methods=['GET', 'POST'])
@login_required
def create_analysis():
    logger.debug(f"[/HOME] Language: {get_locale()}")
    logger.info("[/HOME] Rendering create analysis template...")
    return render_template('analysis/create_analysis.html')

@bp.route('/create_analysis/v0', methods=['GET'])
def create_analysis_v0():
    logger.debug(f"[/V0] Language: {get_locale()}")
    logger.info("[/V0] Rendering create analysis v0 template...")
    return render_template('analysis/create_analysis_v0.html')

@bp.route('/analyze', methods=['GET', 'POST'])
@login_required
def analyze():
    logger.info(f"[/ANALYSIS/ANALYZE] Request to analysis/analyze from {request.remote_addr} with method {request.method}")

    # Leer datos enviados
    text = request.form.get('text', '')
    model = request.form.get('model', '')
    title = request.form.get('title', '')
    authors = request.form.get('authors', '')
    url = request.form.get('url', '')
    analysis_mode = request.form.get('analysis_mode', 'automatic')

    app = bp.get_app() if hasattr(bp, 'get_app') else None

    # Preparar JSON
    payload = {'text': text, 
               'model': model,
               'title': title,
               'authors': authors,
               'url': url,
               'analysis_mode': analysis_mode}

    # Add user information to headers
    headers = API_HEADERS.copy()
    user = get_user_info()
    if user:
        headers['X-User-ID'] = user.get('id', 'anonymous')
        headers['X-User-Email'] = user.get('email', 'anonymous@example.com')
    
    # Llamada a la API
    logger.info(f"[/ANALYSIS/ANALYZE] Sending request to API ({URL_API_ENDPOINT_ANALYSIS_ANALYZE})...")
    resp = requests.post(URL_API_ENDPOINT_ANALYSIS_ANALYZE, json=payload, headers=headers)
    if resp.status_code == 200:
        data = resp.json()
        logger.info(f"[/ANALYSIS/ANALYZE] Received response from API: {data}")
        
        # Save analysis to database
        try:
            analysis_doc = create_analysis_document(data, text, title, authors, url, model, analysis_mode)
            result = db.DB_ANALYSIS.insert_one(analysis_doc)
            analysis_id = str(result.inserted_id)
            logger.info(f"Automatic analysis saved with ID: {analysis_id}")
            
            # Add analysis_id to data for frontend
            data['local_analysis_id'] = analysis_id
            
        except Exception as e:
            logger.error(f"Error saving automatic analysis: {str(e)}")
            # Continue with rendering even if save fails
        
        logger.info("[/ANALYSIS/ANALYZE] Rendering analysis template...")
        return render_template(
            'analysis/analysis.html',
            language=get_locale(),
            data=data,
            highlight_map=HIGHLIGHT_COLOR_MAP,
            contenido_general_variables=CONTENIDO_GENERAL_VARIABLES,
            lenguaje_variables=LENGUAJE_VARIABLES,
            fuentes_variables=FUENTES_VARIABLES,
            api_url_edit=URL_API_ENDPOINT_ANALYSIS_EDITS,
            api_url_save_annotations=URL_API_ENDPOINT_ANALYSIS_SAVE_ANNOTATIONS
        )
    else:
        logger.error("[/ANALYSIS/ANALYZE] Error in API request")
        return jsonify({"error": "Error en la solicitud al API"}), 500


@bp.route('/analyze_manual', methods=['GET', 'POST'])
@login_required
def analyze_manual():
    logger.info(f"[/ANALYSIS/ANALYZE_MANUAL] Request to analysis/analyze_manual from {request.remote_addr} with method {request.method}")

    if request.method == 'GET':
        # Render manual analysis page
        return render_template(
            'analysis/manual_analysis.html',
            language=get_locale(),
            highlight_map=HIGHLIGHT_COLOR_MAP,
            contenido_general_variables=CONTENIDO_GENERAL_VARIABLES,
            lenguaje_variables=LENGUAJE_VARIABLES,
            fuentes_variables=FUENTES_VARIABLES,
            config=config
        )
    
    # POST method - process manual analysis
    if request.is_json:
        # Handle JSON data from frontend
        data = request.get_json()
        text = data.get('text', '')
        title = data.get('title', '')
        authors = data.get('authors', '')
        url = data.get('url', '')
        annotations = data.get('annotations', [])
        selected_topic = data.get('selected_topic', '')
        protagonist_analysis = data.get('protagonist_analysis', {})
        timestamp = data.get('timestamp', datetime.now().isoformat())
    else:
        # Handle form data
        text = request.form.get('text', '')
        title = request.form.get('title', '')
        authors = request.form.get('authors', '')
        url = request.form.get('url', '')
        annotations = []
        selected_topic = ''
        protagonist_analysis = {}
        timestamp = datetime.now().isoformat()

    if not text and not url:
        return jsonify({"error": "Se requiere texto o URL para el análisis"}), 400

    # If we have annotations, save the analysis
    if annotations:
        try:
            # Create analysis document
            analysis_doc = {
                'user_id': session.get('user_id'),
                'text': text,
                'title': title,
                'authors': authors,
                'url': url,
                'analysis_mode': 'manual',
                'status': 'completed',
                'annotations': annotations,
                'selected_topic': selected_topic,
                'protagonist_analysis': protagonist_analysis,
                'created_at': timestamp,
                'updated_at': timestamp,
                'metadata': {
                    'total_annotations': len(annotations),
                    'categories': list(set([ann.get('category', '') for ann in annotations])),
                    'variables': list(set([ann.get('variable', '') for ann in annotations]))
                }
            }
            
            # Save to database
            result = db.DB_ANALYSIS.insert_one(analysis_doc)
            analysis_id = str(result.inserted_id)
            
            logger.info(f"Manual analysis saved with ID: {analysis_id}")
            
            return jsonify({
                "success": True,
                "analysis_id": analysis_id,
                "message": "Análisis guardado correctamente"
            })
            
        except Exception as e:
            logger.error(f"Error saving manual analysis: {str(e)}")
            return jsonify({"error": "Error al guardar el análisis"}), 500

    # For initial manual analysis setup
    manual_data = {
        'text': text,
        'title': title,
        'authors': authors,
        'url': url,
        'analysis_mode': 'manual',
        'status': 'ready_for_manual_analysis'
    }

    return render_template(
        'analysis/manual_analysis.html',
        data=manual_data,
        language=get_locale(),
        highlight_map=HIGHLIGHT_COLOR_MAP,
        contenido_general_variables=CONTENIDO_GENERAL_VARIABLES,
        lenguaje_variables=LENGUAJE_VARIABLES,
        fuentes_variables=FUENTES_VARIABLES,
        config=config
    )


@bp.route('/analyze/v0', methods=['GET', 'POST'])
@login_required
def analyze_v0():

    logger.info(f"[/ANALYSIS/ANALYZE/V0] Request to analysis/analyze/v0 from {request.remote_addr} with method {request.method}")

    # Leer datos enviados
    text = "La ingeniera Laura Gómez fue elegida como la responsable del ambicioso proyecto de inteligencia artificial en la empresa tecnológica Neuronix. Esta decisión marcó un hito importante, ya que es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres. Laura, de 34 años, cuenta con una trayectoria profesional sólida: obtuvo su doctorado en aprendizaje automático en la Universidad de Stanford, trabajó durante cinco años en una reconocida startup de Silicon Valley y publicó múltiples artículos en conferencias internacionales. A pesar de estos logros, algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”. “Me siento honrada y lista para afrontar este reto con todo mi equipo”, declaró en la rueda de prensa. Sin embargo, el jefe del área técnica, Andrés Morales, comentó que “es un cambio arriesgado” y que “espera que el equipo esté preparado para ajustarse a su estilo de liderazgo”. En redes sociales, las reacciones fueron mixtas. Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa. “No es cuestión de género, es cuestión de experiencia real”, escribió un usuario. Paradójicamente, Laura tiene más experiencia que varios de los anteriores líderes del mismo equipo. Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología. Aunque Laura ha demostrado sobradamente su preparación, el foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional."
    model = "avanzado"

    app = bp.get_app() if hasattr(bp, 'get_app') else None

    # Preparar JSON
    payload = {'text': text, 'model': model}

    # Llamada a la API
    print(f"[/ANALYZE/V0] Retrieving default example...")
    data = {'_id': '683977562b6934310b5f8d5c', 'analysis': {'edited': {'contenido_general': None, 'fuentes': None, 'lenguaje': None}, 'original': {'contenido_general': {'genero_periodista': 5, 'genero_personas_mencionadas': [2, 1], 'personas_mencionadas': ['Laura', 'Andrés Morales'], 'tema': 12}, 'fuentes': {'fuentes': [{'declaracion_fuente': 'Me siento honrada y lista para afrontar este reto con todo mi equipo', 'genero_fuente': 2, 'nombre_fuente': 'Laura', 'tipo_fuente': 1}, {'declaracion_fuente': 'es un cambio arriesgado', 'genero_fuente': 1, 'nombre_fuente': 'Andrés Morales', 'tipo_fuente': 1}, {'declaracion_fuente': 'espera que el equipo esté preparado para ajustarse a su estilo de liderazgo', 'genero_fuente': 1, 'nombre_fuente': 'Andrés Morales', 'tipo_fuente': 1}, {'declaracion_fuente': 'No es cuestión de género, es cuestión de experiencia real', 'genero_fuente': 3, 'nombre_fuente': 'usuario anónimo', 'tipo_fuente': 2}]}, 'lenguaje': {'androcentrismo': {'ejemplos_articulo': ['Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología.'], 'etiqueta': [1]}, 'asimetria': {'ejemplos_articulo': ['Algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”.', 'El foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.'], 'etiqueta': [1]}, 'cargos_mujeres': {'ejemplos_articulo': ['Es la primera vez que una mujer lidera un equipo técnico de esa magnitud.'], 'etiqueta': [1]}, 'comparacion_mujeres_hombres': {'ejemplos_articulo': ['Algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”.', 'A pesar de estos logros, algunos medios destacaron su “juventud”.'], 'etiqueta': [1]}, 'denominacion_dependiente': {'ejemplos_articulo': ['El foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.'], 'etiqueta': [1]}, 'denominacion_redundante': {'ejemplos_articulo': ['Un puesto que tradicionalmente había sido ocupado exclusivamente por hombres.'], 'etiqueta': [1]}, 'denominacion_sexualizada': {'ejemplos_articulo': ['El foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.'], 'etiqueta': [1]}, 'dual_aparente': {'ejemplos_articulo': ['La primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres.'], 'etiqueta': [1]}, 'excepcion_noticiabilidad': {'ejemplos_articulo': ['Es la primera vez que una mujer lidera un equipo técnico de esa magnitud.'], 'etiqueta': [1]}, 'hombre_humanidad': {'ejemplos_articulo': ['Puesto que tradicionalmente había sido ocupado exclusivamente por hombres.'], 'etiqueta': [1]}, 'infantilizacion': {'ejemplos_articulo': ['Se refirieron a ella como “una joven promesa con carisma y ambición”.'], 'etiqueta': [1]}, 'lenguaje_sexista': {'ejemplos_articulo': ['Es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres.', 'Algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”.', 'El foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.'], 'etiqueta': [1]}, 'masculino_generico': {'ejemplos_articulo': ['Puesto que tradicionalmente había sido ocupado exclusivamente por hombres.', 'Es un cambio arriesgado y que espera que el equipo esté preparado para ajustarse a su estilo de liderazgo.'], 'etiqueta': [1]}, 'sexismo_social': {'ejemplos_articulo': ['Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa.'], 'etiqueta': [1]}}}}, 'authors': '', 'highlight': {'edited': {'contenido_general': None, 'fuentes': None, 'lenguaje': None}, 'original': {'contenido_general': 'Esta decisión marcó un hito importante, ya que es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres. <mark class="color-3">Laura</mark>, de 34 años, cuenta con una trayectoria profesional sólida: obtuvo su doctorado en aprendizaje automático en la Universidad de Stanford, trabajó durante cinco años en una reconocida startup de Silicon Valley y publicó múltiples artículos en conferencias internacionales. A pesar de estos logros, algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”. “Me siento honrada y lista para afrontar este reto con todo mi equipo”, declaró en la rueda de prensa. Sin embargo, el jefe del área técnica, <mark class="color-3">Andrés Morales</mark>, comentó que “es un cambio arriesgado” y que “espera que el equipo esté preparado para ajustarse a su estilo de liderazgo”. En redes sociales, las reacciones fueron mixtas. Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa. “No es cuestión de género, es cuestión de experiencia real”, escribió un usuario. Paradójicamente, <mark class="color-3">Laura</mark> tiene más experiencia que varios de los anteriores líderes del mismo equipo. Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología. Aunque <mark class="color-3">Laura</mark> ha demostrado sobradamente su preparación, el foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.\r\n', 'fuentes': 'Esta decisión marcó un hito importante, ya que es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres. <mark class="color-3">Laura</mark>, de 34 años, cuenta con una trayectoria profesional sólida: obtuvo su doctorado en aprendizaje automático en la Universidad de Stanford, trabajó durante cinco años en una reconocida startup de Silicon Valley y publicó múltiples artículos en conferencias internacionales. A pesar de estos logros, algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”. “<mark class="color-4">Me siento honrada y lista para afrontar este reto con todo mi equipo</mark>”, declaró en la rueda de prensa. Sin embargo, el jefe del área técnica, <mark class="color-3"><mark class="color-3">Andrés Morales</mark></mark>, comentó que “<mark class="color-4">es un cambio arriesgado</mark>” y que “<mark class="color-4">espera que el equipo esté preparado para ajustarse a su estilo de liderazgo</mark>”. En redes sociales, las reacciones fueron mixtas. Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa. “<mark class="color-4">No es cuestión de género, es cuestión de experiencia real</mark>”, escribió un usuario. Paradójicamente, <mark class="color-3">Laura</mark> tiene más experiencia que varios de los anteriores líderes del mismo equipo. Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología. Aunque <mark class="color-3">Laura</mark> ha demostrado sobradamente su preparación, el foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.\r\n', 'lenguaje': 'Esta decisión marcó un hito importante, ya que es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres. Laura, de 34 años, cuenta con una trayectoria profesional sólida: obtuvo su doctorado en aprendizaje automático en la Universidad de Stanford, trabajó durante cinco años en una reconocida startup de Silicon Valley y publicó múltiples artículos en conferencias internacionales. A pesar de estos logros, algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”. “Me siento honrada y lista para afrontar este reto con todo mi equipo”, declaró en la rueda de prensa. Sin embargo, el jefe del área técnica, Andrés Morales, comentó que “es un cambio arriesgado” y que “espera que el equipo esté preparado para ajustarse a su estilo de liderazgo”. En redes sociales, las reacciones fueron mixtas. <mark class="color-8">Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa.</mark> “No es cuestión de género, es cuestión de experiencia real”, escribió un usuario. Paradójicamente, Laura tiene más experiencia que varios de los anteriores líderes del mismo equipo. <mark class="color-9">Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología.</mark> Aunque Laura ha demostrado sobradamente su preparación, el foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.\r\n'}}, 'model': 'avanzado', 'status': 'ok', 'text': 'Esta decisión marcó un hito importante, ya que es la primera vez que una mujer lidera un equipo técnico de esa magnitud dentro de la compañía, un puesto que tradicionalmente había sido ocupado exclusivamente por hombres. Laura, de 34 años, cuenta con una trayectoria profesional sólida: obtuvo su doctorado en aprendizaje automático en la Universidad de Stanford, trabajó durante cinco años en una reconocida startup de Silicon Valley y publicó múltiples artículos en conferencias internacionales. A pesar de estos logros, algunos medios destacaron su “juventud” y se refirieron a ella como “una joven promesa con carisma y ambición”. “Me siento honrada y lista para afrontar este reto con todo mi equipo”, declaró en la rueda de prensa. Sin embargo, el jefe del área técnica, Andrés Morales, comentó que “es un cambio arriesgado” y que “espera que el equipo esté preparado para ajustarse a su estilo de liderazgo”. En redes sociales, las reacciones fueron mixtas. Mientras muchas personas celebraron el nombramiento como un avance hacia la igualdad de género en el sector tecnológico, otros cuestionaron si su elección respondía a criterios de mérito o simplemente a una estrategia de imagen por parte de la empresa. “No es cuestión de género, es cuestión de experiencia real”, escribió un usuario. Paradójicamente, Laura tiene más experiencia que varios de los anteriores líderes del mismo equipo. Este caso ha abierto nuevamente el debate sobre el tratamiento que reciben las mujeres en puestos de alta responsabilidad, especialmente en sectores tradicionalmente masculinizados como la tecnología. Aunque Laura ha demostrado sobradamente su preparación, el foco mediático sigue centrado más en su condición de mujer que en su capacidad profesional.\r\n', 'timestamp': '2025-05-30T09:16:06.513754+00:00', 'title': '', 'url': ''}

    logger.info("[/ANALYSIS/ANALYZE/V0] Rendering analysis template...")
    return render_template(
        'analysis/analysis.html',
        language=get_locale(),
        data=data,
        highlight_map=HIGHLIGHT_COLOR_MAP,
        contenido_general_variables=CONTENIDO_GENERAL_VARIABLES,
        lenguaje_variables=LENGUAJE_VARIABLES,
        fuentes_variables=FUENTES_VARIABLES,
        api_url_edit=URL_API_ENDPOINT_ANALYSIS_EDITS,
        api_url_save_annotations=URL_API_ENDPOINT_ANALYSIS_SAVE_ANNOTATIONS
    )


@bp.route('/history', methods=['GET'])
@login_required
def analysis_history():
    """Display analysis history page with search and filtering capabilities"""
    logger.info(f"[/ANALYSIS/HISTORY] Request from {request.remote_addr} [{request.method}]")
    
    # Get search and filter parameters
    search_query = request.args.get('search', '').strip()
    date_from = request.args.get('date_from', '').strip()
    date_to = request.args.get('date_to', '').strip()
    model_filter = request.args.get('model', '').strip()
    
    # Prepare API call to get analysis history
    api_url = f"http://{API_HOST}:{API_PORT}/analysis/history"
    
    # Prepare query parameters for API
    api_params = {}
    if search_query:
        api_params['search'] = search_query
    if date_from:
        api_params['date_from'] = date_from
    if date_to:
        api_params['date_to'] = date_to
    if model_filter:
        api_params['model'] = model_filter
    
    # Add user information to headers
    headers = API_HEADERS.copy()
    user = get_user_info()
    if user:
        headers['X-User-ID'] = user.get('id', 'anonymous')
        headers['X-User-Email'] = user.get('email', 'anonymous@example.com')
    
    try:
        # Call API to get analysis history
        response = requests.get(api_url, params=api_params, headers=headers, timeout=10)
        if response.status_code == 200:
            data = response.json()
            analyses = data.get('analyses', [])
        else:
            # Fallback: return empty list if API is not available yet
            analyses = []
            logger.warning(f"API returned status {response.status_code}, using empty list")
    except Exception as e:
        # Fallback: return empty list if API call fails
        analyses = []
        logger.warning(f"Failed to fetch analysis history: {e}")
    
    # Get unique models for filter dropdown
    models = list(set([analysis.get('model', '') for analysis in analyses if analysis.get('model')]))
    models.sort()
    
    return render_template(
        'analysis/analysis_history.html',
        analyses=analyses,
        search_query=search_query,
        date_from=date_from,
        date_to=date_to,
        model_filter=model_filter,
        models=models,
        total_count=len(analyses)
    )


@bp.route('/view/<analysis_id>', methods=['GET'])
@login_required
def view_analysis(analysis_id):
    """Display a specific analysis"""
    logger.info(f"[/ANALYSIS/VIEW/{analysis_id}] Request from {request.remote_addr} [{request.method}]")
    
    try:
        # Call API to get the specific analysis
        api_url = f"http://{API_HOST}:{API_PORT}/analysis/{analysis_id}"
        
        # Add user information to headers
        headers = API_HEADERS.copy()
        user = get_user_info()
        if user:
            headers['X-User-ID'] = user.get('id', 'anonymous')
            headers['X-User-Email'] = user.get('email', 'anonymous@example.com')
        
        response = requests.get(api_url, headers=headers, timeout=10)
        
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                analysis = data.get('analysis')
                
                return render_template(
                    'analysis/analysis.html',
                    language=get_locale(),
                    data=analysis,
                    highlight_map=HIGHLIGHT_COLOR_MAP,
                    contenido_general_variables=CONTENIDO_GENERAL_VARIABLES,
                    lenguaje_variables=LENGUAJE_VARIABLES,
                    fuentes_variables=FUENTES_VARIABLES,
                    api_url_edit=URL_API_ENDPOINT_ANALYSIS_EDITS,
                    api_url_save_annotations=URL_API_ENDPOINT_ANALYSIS_SAVE_ANNOTATIONS
                )
            else:
                abort(404, description=data.get('error', 'Análisis no encontrado'))
        else:
            abort(response.status_code, description="Error al obtener el análisis")
            
    except Exception as e:
        logger.error(f"Error viewing analysis {analysis_id}: {e}")
        abort(500, description="Error interno del servidor")
