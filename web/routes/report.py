# web/report.py
from flask import Blueprint, render_template, request, jsonify, session, send_file, abort
import requests
from flask_babel import get_locale
import configparser
import ast
from io import BytesIO
from datetime import datetime
from collections import Counter, OrderedDict
from web.utils.logger import logger
from web.utils.decorators import challenge_restricted


# ----------------- BLUEPRINT -----------------
bp = Blueprint(
    'report',
    __name__,
    url_prefix='/report'
)

# ----------------- CONFIG -----------------
config = configparser.ConfigParser()
config.read('config.ini')

WEB_HOST = config['WEB']['HOST']
WEB_PORT = config['WEB']['PORT']
DEBUG = config['WEB'].getboolean('DEBUG')
API_HOST = config['API']['HOST']
API_PORT = config['API']['PORT']
API_HEADERS_str = config['API']['HEADERS']
API_HEADERS = ast.literal_eval(API_HEADERS_str)
ENDPOINT_ANALYSIS = config['API']['ENDPOINT_ANALYSIS']
ENDPOINT_ANALYSIS_ANALYZE = config['API']['ENDPOINT_ANALYSIS_ANALYZE']
ENDPOINT_DATA = config['API']['ENDPOINT_DATA']
ENDPOINT_DATA_GET_CONTEXTO = config['API']['ENDPOINT_DATA_GET_CONTEXTO']
COLLECTION_DATA = config['DATABASE']['COLLECTION_DATA']
OPENAI_API_KEY = config['API'].get('OPENAI_API_KEY', '')
OPENAI_MODEL = config['API'].get('OPENAI_MODEL', 'gpt-4o-mini') or 'gpt-4o-mini'

# ----------------- URLs -----------------
URL_API_ENDPOINT_ANALYSIS_ANALYZE = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_ANALYSIS}/{ENDPOINT_ANALYSIS_ANALYZE}"
URL_API_ENDPOINT_DATA = f"http://{API_HOST}:{API_PORT}/{ENDPOINT_DATA}"
URL_API_ENDPOINT_DATA_COLLECTION = f"{URL_API_ENDPOINT_DATA}/{COLLECTION_DATA}"
URL_API_ENDPOINT_DATA_GET_CONTEXTO = f"{URL_API_ENDPOINT_DATA}/{ENDPOINT_DATA_GET_CONTEXTO}"


# ----------------- VARIABLE CATALOG -----------------
# Report-friendly categorical variables → (human label, config map key)
REPORT_VARIABLES = OrderedDict([
    ('menciona_ia',        ('Menciona IA',                    'MENCIONA_IA')),
    ('ia_tema_central',    ('IA como tema central',           'IA_TEMA_CENTRAL')),
    ('significado_ia',     ('Explicación del significado de IA', 'SIGNIFICADO_IA')),
    ('tema',               ('Temática de la noticia',         'TEMA')),
    ('genero_periodista',  ('Género del/la periodista',       'GENERO_PERIODISTA')),
    ('genero_personas',    ('Género de las personas mencionadas', 'GENERO_PERSONAS_MENCIONADAS')),
    ('cita_titular',       ('Cita en el titular',             'CITA_TITULAR')),
    ('lenguaje_sexista',   ('Lenguaje sexista',               'LENGUAJE_SEXISTA')),
])


def _load_map(cfg_key):
    """Parse a mapping dict from config.ini and normalise keys to str."""
    try:
        raw = ast.literal_eval(config['VARIABLES'][cfg_key]) if config.has_option('VARIABLES', cfg_key) \
            else ast.literal_eval(config['DEFAULT'][cfg_key])
    except Exception:
        # Fall back to a flat scan of every section
        raw = {}
        for section in config.sections():
            if config.has_option(section, cfg_key):
                try:
                    raw = ast.literal_eval(config[section][cfg_key])
                    break
                except Exception:
                    continue
    return {str(k): v for k, v in (raw or {}).items()}


def _fetch_contextos():
    try:
        r = requests.get(URL_API_ENDPOINT_DATA_GET_CONTEXTO, timeout=8)
        r.raise_for_status()
        return r.json().get('contextos', [])
    except Exception as e:
        logger.warning(f"[/REPORT] No se pudieron obtener contextos: {e}")
        return []


def _fetch_data():
    r = requests.get(URL_API_ENDPOINT_DATA_COLLECTION, timeout=15)
    r.raise_for_status()
    data = r.json()
    return data if isinstance(data, list) else data.get('data', [])


def _parse_fecha(s):
    for fmt in ('%d/%m/%Y', '%Y-%m-%d'):
        try:
            return datetime.strptime(str(s), fmt).date()
        except Exception:
            continue
    return None


def _filter_rows(rows, context, fd, fh):
    d_from = _parse_fecha(fd) if fd else None
    d_to = _parse_fecha(fh) if fh else None
    out = []
    for it in rows:
        if context and it.get('contexto') != context:
            continue
        if d_from or d_to:
            f = _parse_fecha(it.get('Fecha'))
            if f is None:
                continue
            if d_from and f < d_from:
                continue
            if d_to and f > d_to:
                continue
        out.append(it)
    return out


def _year_of(it):
    f = _parse_fecha(it.get('Fecha'))
    return f.year if f else None


def _build_tables(rows, variables):
    """For each variable, a contingency table (value x year) with % col. and % fila."""
    tables = []
    for field in variables:
        if field not in REPORT_VARIABLES:
            continue
        label, cfg_key = REPORT_VARIABLES[field]
        value_map = _load_map(cfg_key)

        freq = {}          # value -> {year -> count}
        years_set = set()
        for it in rows:
            v = it.get(field)
            if v is None or str(v).strip() in ('', 'None', 'nan', '-1'):
                continue
            y = _year_of(it)
            if y is None:
                continue
            v = str(v)
            freq.setdefault(v, {})
            freq[v][y] = freq[v].get(y, 0) + 1
            years_set.add(y)

        years = sorted(years_set)
        col_totals = {y: sum(freq[v].get(y, 0) for v in freq) for y in years}
        row_totals = {v: sum(freq[v].values()) for v in freq}
        grand = sum(row_totals.values())
        if grand == 0:
            continue

        ordered = sorted(freq.keys(), key=lambda v: row_totals[v], reverse=True)
        data_rows = []
        for v in ordered:
            cells = []
            for y in years:
                fr = freq[v].get(y, 0)
                pc = (fr / col_totals[y] * 100) if col_totals[y] else 0
                pf = (fr / row_totals[v] * 100) if row_totals[v] else 0
                cells.append({'pct_col': f"{pc:.2f}", 'pct_fila': f"{pf:.2f}"})
            data_rows.append({
                'label': value_map.get(v, f"Código {v}"),
                'cells': cells,
                'total_pct_col': f"{(row_totals[v] / grand * 100):.2f}",
                'total_frec': row_totals[v],
            })

        total_cells = [{'pct_col': '100', 'pct_fila': f"{(col_totals[y] / grand * 100):.2f}"} for y in years]

        tables.append({
            'field': field,
            'title': label,
            'years': [str(y) for y in years],
            'rows': data_rows,
            'total_cells': total_cells,
            'grand': grand,
        })
    return tables


def _table_to_text(tbl):
    """Plain-text rendering of a contingency table for the LLM prompt."""
    header = [tbl['title']]
    for y in tbl['years']:
        header += [f"{y} %col.", f"{y} %fila"]
    header += ["Total %col.", "Total Frec."]
    lines = [" | ".join(header)]
    for r in tbl['rows']:
        row = [str(r['label'])]
        for c in r['cells']:
            row += [f"{c['pct_col']}%", f"{c['pct_fila']}%"]
        row += [f"{r['total_pct_col']}%", str(r['total_frec'])]
        lines.append(" | ".join(row))
    trow = ["Total"]
    for c in tbl['total_cells']:
        trow += [f"{c['pct_col']}%", f"{c['pct_fila']}%"]
    trow += ["100%", str(tbl['grand'])]
    lines.append(" | ".join(trow))
    return "\n".join(lines)


def _ai_narrative(tbl, context, periodo):
    """1–2 paragraph descriptive analysis of a table via OpenAI (optional)."""
    if not OPENAI_API_KEY:
        return None
    try:
        from openai import OpenAI
    except Exception as e:
        logger.warning(f"[/REPORT] openai no disponible: {e}")
        return None
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        table_txt = _table_to_text(tbl)
        prompt = (
            "Actúa como un experto sociólogo y analista de datos con perspectiva de género. "
            "Te paso una tabla de contingencia de una variable por año, con porcentajes por columna (% col.) "
            "y por fila (% fila). Escribe 1 o 2 párrafos concisos y profesionales, en español, analizando los "
            "resultados más destacables para un informe formal. Responde directamente, sin preámbulos ni títulos.\n\n"
            f"Variable: {tbl['title']}\nBase de datos: {context}\nPeriodo: {periodo}\n\nTabla:\n{table_txt}"
        )
        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        logger.warning(f"[/REPORT] Error IA narrative: {e}")
        return None


def _prepare_report(args):
    context = (args.get('context') or '').strip()
    fd = args.get('fecha_desde') or ''
    fh = args.get('fecha_hasta') or ''
    include_ai = (args.get('ai') == '1')
    variables = args.getlist('vars') if hasattr(args, 'getlist') else args.get('vars', [])
    variables = [v for v in variables if v in REPORT_VARIABLES]

    rows = _fetch_data()
    filtered = _filter_rows(rows, context, fd, fh)
    tables = _build_tables(filtered, variables)

    meta = {
        'context': context or 'Todas las bases de datos',
        'fecha_desde': fd,
        'fecha_hasta': fh,
        'total': len(filtered),
        'variables': variables,
        'include_ai': include_ai,
    }

    if include_ai:
        periodo = f"{fd or 'inicio'} — {fh or 'hoy'}"
        for t in tables:
            t['analysis'] = _ai_narrative(t, meta['context'], periodo)

    return meta, tables


#  ----------------- ENDPOINTS -----------------
@bp.route('/create', methods=['GET', 'POST'])
@challenge_restricted
def create():
    logger.info(f"[/REPORT/CREATE] Request to report/create from {request.remote_addr} with method {request.method}")
    return render_template('report/create_report.html')


@bp.route('/generate', methods=['GET'])
@challenge_restricted
def generate():
    logger.info(f"[/REPORT/GENERATE] Request from {request.remote_addr}")
    contextos = _fetch_contextos()

    meta, tables = None, None
    submitted = bool(request.args.get('vars'))
    if submitted:
        try:
            meta, tables = _prepare_report(request.args)
        except Exception as e:
            logger.error(f"[/REPORT/GENERATE] Error building preview: {e}")
            abort(502, description=f"No se pudo generar la vista previa: {e}")

    return render_template(
        'report/generate_report.html',
        contextos=contextos,
        variables=REPORT_VARIABLES,
        meta=meta,
        tables=tables,
        selected_context=(request.args.get('context') or '').strip(),
        selected_vars=request.args.getlist('vars'),
        fecha_desde=request.args.get('fecha_desde', ''),
        fecha_hasta=request.args.get('fecha_hasta', ''),
    )


@bp.route('/download', methods=['GET'])
@challenge_restricted
def download():
    logger.info(f"[/REPORT/DOWNLOAD] Request from {request.remote_addr}")
    if not request.args.get('vars'):
        abort(400, description="Selecciona al menos una variable.")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        abort(500, description=f"python-docx no disponible: {e}")

    meta, tables = _prepare_report(request.args)

    doc = Document()
    doc.add_heading('Informe IRIS', 0)

    sub = doc.add_paragraph()
    sub.add_run('Base de datos: ').bold = True
    sub.add_run(meta['context'])
    periodo = doc.add_paragraph()
    periodo.add_run('Periodo: ').bold = True
    periodo.add_run(f"{meta['fecha_desde'] or 'inicio'} — {meta['fecha_hasta'] or 'hoy'}")
    ntot = doc.add_paragraph()
    ntot.add_run('Total de noticias analizadas: ').bold = True
    ntot.add_run(str(meta['total']))
    doc.add_paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    for tbl in tables:
        doc.add_heading(tbl['title'], level=1)
        years = tbl['years']
        ncols = 1 + 2 * len(years) + 2
        t = doc.add_table(rows=2, cols=ncols)
        t.style = 'Table Grid'

        # Header row 1: variable name + year group labels (merged 2 cells) + Total
        h1 = t.rows[0].cells
        h1[0].text = tbl['title']
        h1[0].merge(t.rows[1].cells[0])
        idx = 1
        for y in years + ['Total']:
            left = h1[idx]
            left.merge(h1[idx + 1])
            left.text = y
            idx += 2

        # Header row 2: % col. / % fila (and Total: % col. / Frec.)
        h2 = t.rows[1].cells
        idx = 1
        for _ in years:
            h2[idx].text = '% col.'
            h2[idx + 1].text = '% fila'
            idx += 2
        h2[idx].text = '% col.'
        h2[idx + 1].text = 'Frec.'

        # Data rows
        for r in tbl['rows']:
            cells = t.add_row().cells
            cells[0].text = str(r['label'])
            idx = 1
            for c in r['cells']:
                cells[idx].text = f"{c['pct_col']}%"
                cells[idx + 1].text = f"{c['pct_fila']}%"
                idx += 2
            cells[idx].text = f"{r['total_pct_col']}%"
            cells[idx + 1].text = str(r['total_frec'])

        # Total row
        tr = t.add_row().cells
        tr[0].text = 'Total'
        idx = 1
        for c in tbl['total_cells']:
            tr[idx].text = f"{c['pct_col']}%"
            tr[idx + 1].text = f"{c['pct_fila']}%"
            idx += 2
        tr[idx].text = '100%'
        tr[idx + 1].text = str(tbl['grand'])

        doc.add_paragraph()

        if tbl.get('analysis'):
            doc.add_heading('Análisis descriptivo', level=2)
            doc.add_paragraph(tbl['analysis'])
            doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"informe_iris_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=fname,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
