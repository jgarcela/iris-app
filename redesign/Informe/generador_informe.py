import pandas as pd
import ollama
from docx import Document
import os

def generar_tabla(df, col_fila, col_columna, mapeo_fila=None, mapeo_col=None):
    # Crear crosstab con frecuencias absolutas
    crosstab_frec = pd.crosstab(df[col_fila], df[col_columna])
    if mapeo_fila:
        crosstab_frec.index = crosstab_frec.index.map(lambda x: mapeo_fila.get(x, x))
    if mapeo_col:
        crosstab_frec.columns = crosstab_frec.columns.map(lambda x: mapeo_col.get(x, x))
        
    # Variables de totales
    col_totals = crosstab_frec.sum(axis=0)
    row_totals = crosstab_frec.sum(axis=1)
    total_general = col_totals.sum()
    
    # Configuramos las cabeceras jerárquicas
    headers_nivel1 = []
    headers_nivel2 = []
    
    for c in crosstab_frec.columns:
        headers_nivel1.extend([c, c])
        headers_nivel2.extend(["% col.", "% fila"])
        
    headers_nivel1.extend(["Total", "Total"])
    headers_nivel2.extend(["% col.", "Frec."])
    
    multi_idx = pd.MultiIndex.from_arrays([headers_nivel1, headers_nivel2])
    df_output = pd.DataFrame(index=crosstab_frec.index, columns=multi_idx)

    # Rellenar datos iterativamente
    for c in crosstab_frec.columns:
        # Porcentaje columna
        df_output[(c, "% col.")] = (crosstab_frec[c] / col_totals[c] * 100).apply(lambda x: f"{x:.2f}%")
        # Porcentaje fila
        df_output[(c, "% fila")] = (crosstab_frec[c] / row_totals * 100).apply(lambda x: f"{x:.2f}%")
        
    # Columna Total
    df_output[("Total", "% col.")] = (row_totals / total_general * 100).apply(lambda x: f"{x:.2f}%")
    df_output[("Total", "Frec.")] = row_totals
    
    # Fila "Total"
    df_output.loc["Total"] = ""
    for c in crosstab_frec.columns:
        df_output.loc["Total", (c, "% col.")] = "100%"
        df_output.loc["Total", (c, "% fila")] = f"{(col_totals[c] / total_general * 100):.2f}%"
    df_output.loc["Total", ("Total", "% col.")] = "100%"
    df_output.loc["Total", ("Total", "Frec.")] = total_general
    
    # Renombrar index (para imitar "7. Nombre propio en título")
    df_output.index.name = col_fila
    
    return df_output

def obtener_redaccion_ollama(df_str, modelo):
    prompt = f"""
Actúa como un experto sociólogo y analista de datos. Te voy a pasar una tabla de contingencia con frecuencias absolutas y porcentajes (por columna y por fila).
Tu objetivo es escribir 1 o 2 párrafos redactados analizando los resultados más destacables de manera concisa y profesional, ideal para un informe formal.
Por favor, responde directamente con la redacción en español sin preámbulos.

Tabla:
{df_str}
    """
    try:
        response = ollama.chat(model=modelo, messages=[{'role': 'user', 'content': prompt}])
        return response['message']['content']
    except Exception as e:
        return f"Error al generar texto con Ollama ({modelo}): {str(e)}"

def main():
    # Caminos actualizados para la estructura Experimentos/Informe/
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    csv_path = os.path.join(base_dir, "data", "2026_02_10_imio_def_todo_envio_heidy.xlsx - 2026_02_09_imio_def_todo_clara_scrape.csv")
    
    if not os.path.exists(csv_path):
        print("El archivo CSV no existe en la ruta esperada:", csv_path)
        return
        
    print("Cargando datos...")
    df = pd.read_csv(csv_path)
    
    variable_columna = "año_agrupado"
    mapeo_col = {
        1: "de 2017 a 2021",
        2: "2024"
    }

    columnas_excluidas = [
        'no_N_noticia', 'IdNoticia', 'Medio_num', 'Fecha', 'año', 'año_agrupado',
        'Caracteres', 'Titular', 'no_NombreUsuario', 'no_Autor', 'no_MES',
        'no_Contenido', 'no_Pagina_url', 'no_verificacion', 'no_textonoticia',
        'no_vacio_numero_caracteres', 'contenido_articulo', 'nombre_periodista'
    ]
    
    variables_a_cruzar = []
    for col in df.columns:
        # Excluimos si está en la lista de ignoradas O si tiene más de 30 categorías distintas (evita volver loco o colgar el LLM)
        if col not in columnas_excluidas and df[col].nunique() <= 30:
            variables_a_cruzar.append({
                "variable_fila": col,
                # Formatear el nombre de la variable para que quede bonito como título
                "titulo": col.replace("_", " ").capitalize()
            })
    
    modelo_elegido = "gemma4:e4b" 
    
    # Exportar a Word general
    doc_path = os.path.join(os.path.dirname(__file__), "informe_completo.docx")
    doc = Document()
    doc.add_heading("Resultados del Informe Completo", 0)
    
    # Preparar Excel general
    excel_path = os.path.join(os.path.dirname(__file__), "todas_las_tablas.xlsx")
    excel_writer = pd.ExcelWriter(excel_path, engine='openpyxl')
    
    total_vars = len(variables_a_cruzar)
    for i, var_info in enumerate(variables_a_cruzar):
        var_fila = var_info["variable_fila"]
        mapeo_fila = var_info.get("mapeo_fila", None)  # Opcional
        titulo = var_info.get("titulo", var_fila)
        
        if var_fila not in df.columns:
            print(f"[{i+1}/{total_vars}] La variable {var_fila} no existe en el DataFrame. Saltando...")
            continue
            
        print(f"\n[{i+1}/{total_vars}] Generando tabla para: {var_fila}...")
        df_tabla = generar_tabla(df, var_fila, variable_columna, mapeo_fila, mapeo_col)
        df_tabla.index.name = titulo
        
        # Exportar hoja excel (limitamos el length del nombre de hoja a 31 chars si es necesario)
        sheet_name = var_fila[:31]
        df_tabla.to_excel(excel_writer, sheet_name=sheet_name)
        
        # Generar texto con ollama
        tabla_str = df_tabla.to_markdown()
        print(f"Realizando consulta a Ollama ({modelo_elegido})...")
        texto_ollama = obtener_redaccion_ollama(tabla_str, modelo_elegido)
        
        # Inserción en Word
        doc.add_heading(f"Distribución: {titulo} vs {variable_columna}", level=1)
        
        # Tabla Word
        t = doc.add_table(rows=df_tabla.shape[0] + 2, cols=len(df_tabla.columns) + 1)
        t.style = 'Table Grid'
        
        # Encabezados
        col_names = list(df_tabla.columns)
        t.cell(0, 0).text = df_tabla.index.name or var_fila
        
        for c_idx, c_tuple in enumerate(col_names):
            t.cell(0, c_idx + 1).text = str(c_tuple[0])
            t.cell(1, c_idx + 1).text = str(c_tuple[1])
            
        # Datos
        for r_idx, (index, row) in enumerate(df_tabla.iterrows()):
            t.cell(r_idx + 2, 0).text = str(index)
            for c_idx, val in enumerate(row):
                t.cell(r_idx + 2, c_idx + 1).text = str(val)
                
        doc.add_paragraph()
        doc.add_heading(f"Análisis descriptivo ({var_fila})", level=2)
        doc.add_paragraph(texto_ollama)
        doc.add_page_break()  # Salto de página para que quede más limpio
        
        # Guardar word de manera progresiva por si el programa se interrumpe
        doc.save(doc_path)
        
    excel_writer.close()
    print(f"\nExcel completo guardado en {excel_path}")
    print(f"Word completo actualizado y guardado en {doc_path}")
    print("¡Proceso finalizado con éxito!")

if __name__ == "__main__":
    main()
